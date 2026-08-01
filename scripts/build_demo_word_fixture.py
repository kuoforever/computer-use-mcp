"""Build the polished public-web research fixture used by the GUI Demo."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)


def _font(run, *, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _body(doc: Document, text: str, *, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    _font(paragraph.add_run(text), size=11, color=RGBColor(0, 0, 0))


def build(path: Path) -> None:
    doc = Document()
    settings = doc.settings.element
    view = settings.find(qn("w:view"))
    if view is None:
        view = OxmlElement("w:view")
        settings.append(view)
    view.set(qn("w:val"), "print")
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.append(zoom)
    zoom.set(qn("w:percent"), "100")
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _font(
        header.add_run("DIGITAL WORKPLACE | SOURCE REVIEW"),
        size=9,
        color=GRAY,
        bold=True,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(footer.add_run("Public-source research | Internal Demo"), size=8.5, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)
    _font(
        title.add_run("WORD COLLABORATION RESEARCH NOTES"),
        size=23,
        color=RGBColor(0, 0, 0),
        bold=True,
    )

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _font(
        subtitle.add_run("Microsoft Support public-source review | Working draft"),
        size=14,
        color=GRAY,
    )

    for label, value in (
        ("Date", "July 30, 2026"),
        ("Analyst", "Digital Workplace Research"),
        ("Source", "Microsoft Support - public documentation"),
        ("Status", "Draft notes pending verified browser follow-up"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(f"{label}: "), size=11, color=RGBColor(0, 0, 0), bold=True)
        _font(p.add_run(value), size=11, color=RGBColor(0, 0, 0))

    doc.add_heading("Research question", level=1)
    _body(
        doc,
        "How does Microsoft describe collaborative work in Word, and which "
        "capabilities matter most for a concise internal adoption note?",
    )

    doc.add_heading("Working notes", level=1)
    _body(
        doc,
        "The review uses a public Microsoft Support page in a dedicated browser "
        "profile. No account, tenant, private file, or signed-in session is required.",
    )
    _body(
        doc,
        "The browser step must remain read-only: inspect the page, move through "
        "the article, and collect bounded evidence without posting or downloading.",
    )
    _body(
        doc,
        "The final note should identify the collaboration model and practical "
        "review mechanisms, while keeping the original source visible and reviewable.",
    )

    doc.add_heading("Draft takeaway", level=1)
    _body(
        doc,
        "Treat the shared document as the center of review. Prefer explicit "
        "comments and accountable follow-up over circulating additional copies.",
    )

    doc.add_heading("Verified source follow-up", level=1)
    prompt = doc.add_paragraph()
    prompt.paragraph_format.space_after = Pt(4)
    _font(
        prompt.add_run("Verified browser summary will be added below:"),
        size=11,
        color=DARK_BLUE,
        bold=True,
    )
    _body(doc, "[Pending controlled desktop update]", after=0)

    doc.core_properties.title = "Word Collaboration Research Notes"
    doc.core_properties.subject = "Public-source cross-application Demo fixture"
    doc.core_properties.author = "Guarded Desktop Agent Demo"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("usage: build_demo_word_fixture.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
